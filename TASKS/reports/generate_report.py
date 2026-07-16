"""Generate DOCX reconstruction report for msfs_autoland at baseline 23599b6."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): color_hex,
        qn('w:val'): 'clear',
    })
    shading.append(shd)


def add_table_row(table, cells_data, bold=False, header=False):
    """Add a row to table. cells_data = list of strings."""
    row = table.add_row()
    for i, text in enumerate(cells_data):
        cell = row.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(str(text))
        run.font.size = Pt(9)
        run.font.name = 'Consolas'
        if bold or header:
            run.bold = True
        if header:
            set_cell_shading(cell, '1F3A5F')
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    return row


def create_report():
    doc = Document()

    # Page setup - A4
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    # Default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x1F, 0x1F, 0x1F)

    # ══════════════════════════════════════════════════════════════
    # TITLE
    # ══════════════════════════════════════════════════════════════
    title = doc.add_heading('MSFS AutoLand — Read-Only Reconstruction Report', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Trusted Baseline: 23599b6b4a805b3fe219bbc877b828188dbde221')
    run.font.size = Pt(12)
    run.bold = True

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run('BOOTSTRAP_STATUS: AWAITING_OWNER_REVIEW')
    run2.font.size = Pt(11)
    run2.bold = True
    run2.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)

    doc.add_paragraph()  # spacer

    # ══════════════════════════════════════════════════════════════
    # 1. GIT STATE
    # ══════════════════════════════════════════════════════════════
    doc.add_heading('1. Git State', level=1)

    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header
    hdr = table.rows[0].cells
    for i, text in enumerate(['Check', 'Value', 'Status']):
        hdr[i].text = ''
        p = hdr[i].paragraphs[0]
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(9)
        run.font.name = 'Consolas'
        set_cell_shading(hdr[i], '1F3A5F')
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    rows_data = [
        ['HEAD', '23599b6b4a805b3fe219bbc877b828188dbde221', 'MATCHES baseline'],
        ['master', '23599b6b4a805b3fe219bbc877b828188dbde221', 'MATCHES baseline'],
        ['origin/master', '23599b6b4a805b3fe219bbc877b828188dbde221', 'MATCHES baseline'],
        ['Tree SHA', 'f2a19ec8a77a4efce6b4b55a3fb08822473bc959', 'MATCHES expected'],
        ['Tracked diff', '2 deleted files (tasks/TASK-002, tasks/TASK-006)', 'working tree only'],
        ['Untracked', 'TASKS/bootstrap.md', 'not committed'],
        ['Status', 'working tree not clean (2 D, 1 ??)', '--'],
    ]
    for row_data in rows_data:
        add_table_row(table, row_data)

    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════
    # 2. ENVIRONMENT
    # ══════════════════════════════════════════════════════════════
    doc.add_heading('2. Environment', level=1)

    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, text in enumerate(['Component', 'Value']):
        hdr[i].text = ''
        p = hdr[i].paragraphs[0]
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(9)
        run.font.name = 'Consolas'
        set_cell_shading(hdr[i], '1F3A5F')
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    env_rows = [
        ['Python', '3.14.5'],
        ['pytest', '9.0.3'],
        ['SimConnect', '0.4.26'],
        ['Interpreter', 'C:\\BAT\\venvs\\msfs_autoland_23599_py314\\Scripts\\python.exe'],
    ]
    for row_data in env_rows:
        add_table_row(table, row_data)

    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════
    # 3. TEST SUITE
    # ══════════════════════════════════════════════════════════════
    doc.add_heading('3. Test Suite (208 tests, 0 failed)', level=1)

    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, text in enumerate(['File', 'Count', 'Coverage Scope']):
        hdr[i].text = ''
        p = hdr[i].paragraphs[0]
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(9)
        run.font.name = 'Consolas'
        set_cell_shading(hdr[i], '1F3A5F')
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    test_rows = [
        ['test_safety_guard.py', '28', 'G1-G5 rules, debounce, integration, logging, red-without-fix'],
        ['test_control_ownership.py', '7', 'Channel ownership matrix, no competing commands'],
        ['test_takeover_safety.py', '14', 'Hard safety gates, readback, production readback'],
        ['test_sink_rate_guard.py', '16', 'sink_rate_safe, mode classification (ILS/VOR), timeout'],
        ['test_approach_lifecycle.py', '3', 'State reset on repeat approach'],
        ['test_ils_takeover_crossing.py', '14', 'DH crossing detection, DH guard, fail-closed'],
        ['test_loc_approach.py', '25', 'LOC routing, signal loss, lateral CDI, pipeline, neighbours'],
        ['test_wp0_smoke.py', '4', 'Smoke: fakes, takeover instantiation'],
        ['test_telemetry_recorder.py', '22', 'Schema, immediate write, pending frame, actuator-before-flush'],
        ['test_architecture.py', '9', 'Circular deps, layer separation, DI, forbidden imports'],
        ['test_runway_units.py', '5', 'Feet-to-meters conversion, boundary'],
        ['test_synthetic_glidepath.py', '19', 'MSL/AGL chain, MDA floor/hysteresis, production path'],
        ['replay/test_replay_scenarios.py', '4', 'ILS replay scenarios'],
        ['test_engine_failure.py (root)', '15', 'Asymmetric thrust, engine failure detection'],
    ]
    for row_data in test_rows:
        add_table_row(table, row_data)

    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════
    # 4. ENTRY POINTS
    # ══════════════════════════════════════════════════════════════
    doc.add_heading('4. Entry Points', level=1)

    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, text in enumerate(['File:Line', 'Function/Class', 'Behavior', 'Test Coverage']):
        hdr[i].text = ''
        p = hdr[i].paragraphs[0]
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(9)
        run.font.name = 'Consolas'
        set_cell_shading(hdr[i], '1F3A5F')
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    entry_rows = [
        ['main.py:61', 'AutoLandSystem.__init__', 'Creates all subsystems', 'test_wp0_smoke (instantiation)'],
        ['main.py:108', 'AutoLandSystem.connect()', 'MSFS connection, adapter/optimizer init', 'NOT COVERED'],
        ['main.py:229', 'AutoLandSystem.configure_approach()', 'ILS/VOR/NDB/LOC setup, VAPP calc', 'NOT COVERED'],
        ['main.py:310', 'AutoLandSystem.start_approach()', 'Reset per-approach state, frequencies', 'test_approach_lifecycle (reset)'],
        ['main.py:359', 'AutoLandSystem.stop_approach()', 'Stop, recorder flush', 'test_telemetry_recorder'],
        ['main.py:408', 'AutoLandSystem.execute_go_around()', 'Full throttle, 1500fpm, flaps 2, reset', 'test_telemetry_recorder'],
        ['main.py:525', 'AutoLandSystem.execute_approach()', 'Main loop: telemetry->guard->phase', 'test_telemetry_recorder (2 iter)'],
    ]
    for row_data in entry_rows:
        add_table_row(table, row_data)

    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════
    # 5. SAFETY GUARDS
    # ══════════════════════════════════════════════════════════════
    doc.add_heading('5. Safety Guards', level=1)

    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, text in enumerate(['File:Line', 'Class/Function', 'Behavior', 'Test Coverage']):
        hdr[i].text = ''
        p = hdr[i].paragraphs[0]
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(9)
        run.font.name = 'Consolas'
        set_cell_shading(hdr[i], '1F3A5F')
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    safety_rows = [
        ['safety_guard.py:80', 'ApproachSafetyGuard', 'Pure eval: snapshot->decision. FINAL only.', 'test_safety_guard (28)'],
        ['safety_guard.py:131', 'evaluate() G1', 'abs(vs) > 1500 -> GO_AROUND', 'VERIFIED'],
        ['safety_guard.py:152', 'evaluate() G2', 'bank > 15.0 -> GO_AROUND', 'VERIFIED'],
        ['safety_guard.py:162', 'evaluate() G3', 'airspeed < vref - 10 -> GO_AROUND', 'VERIFIED'],
        ['safety_guard.py:174', 'evaluate() G4', 'airspeed > vref + 20 -> GO_AROUND', 'VERIFIED'],
        ['safety_guard.py:131', 'evaluate() G5', 'Invalid telemetry -> GO_AROUND', 'VERIFIED'],
        ['safety_guard.py:195', 'reset()', 'Clears counters + go_around_executed', 'VERIFIED'],
    ]
    for row_data in safety_rows:
        add_table_row(table, row_data)

    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════
    # 6. CONTROL OWNERSHIP
    # ══════════════════════════════════════════════════════════════
    doc.add_heading('6. Control Ownership (WP-5)', level=1)

    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, text in enumerate(['File:Line', 'Function', 'Behavior', 'Test Coverage']):
        hdr[i].text = ''
        p = hdr[i].paragraphs[0]
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(9)
        run.font.name = 'Consolas'
        set_cell_shading(hdr[i], '1F3A5F')
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    ownership_rows = [
        ['control_ownership.py:26', 'compute_ownership()', 'One owner per channel (roll/pitch/throttle)', 'test_control_ownership (7)'],
        ['control_ownership.py:48', 'GO_AROUND policy', 'NONE/NONE/AP', 'VERIFIED'],
        ['control_ownership.py:55', 'No confirmed takeover', 'AP/AP/AP', 'VERIFIED'],
        ['control_ownership.py:71', 'Confirmed + vJoy ready', 'EXTERNAL/EXTERNAL/throttle_owner', 'VERIFIED'],
    ]
    for row_data in ownership_rows:
        add_table_row(table, row_data)

    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════
    # 7. AUTOPILOT TAKEOVER
    # ══════════════════════════════════════════════════════════════
    doc.add_heading('7. Autopilot Takeover', level=1)

    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, text in enumerate(['File:Line', 'Function', 'Behavior', 'Test Coverage']):
        hdr[i].text = ''
        p = hdr[i].paragraphs[0]
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(9)
        run.font.name = 'Consolas'
        set_cell_shading(hdr[i], '1F3A5F')
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    takeover_rows = [
        ['autopilot_takeover.py:85', 'should_initiate_takeover()', 'ILS: DH+50 window. VOR/NDB/LOC: dist+alt+phase', 'test_ils_takeover_crossing (14)'],
        ['autopilot_takeover.py:140', 'perform_takeover()', 'Timeout->safety->hard/retry->commands->readback', 'test_takeover_safety (14)'],
        ['autopilot_takeover.py:244', '_perform_safety_checks()', 'airborne, attitude, speed, alt, sink_rate', 'test_sink_rate_guard (6)'],
        ['autopilot_takeover.py:313', '_verify_readback()', 'Adapter priority. None=fail-closed', 'test_takeover_safety (4)'],
    ]
    for row_data in takeover_rows:
        add_table_row(table, row_data)

    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════
    # 8. APPROACH PHASES (STATE PATTERN)
    # ══════════════════════════════════════════════════════════════
    doc.add_heading('8. Approach Phases (State Pattern)', level=1)

    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, text in enumerate(['File:Line', 'Class', 'Behavior', 'Test Coverage']):
        hdr[i].text = ''
        p = hdr[i].paragraphs[0]
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(9)
        run.font.name = 'Consolas'
        set_cell_shading(hdr[i], '1F3A5F')
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    phase_rows = [
        ['approach_phases.py:45', 'InitialPhaseState', 'DME check, heading hold, transition to INTERMEDIATE', 'NOT COVERED'],
        ['approach_phases.py:80', 'IntermediatePhaseState', 'Takeover trigger, fix check, transition to FINAL', 'NOT COVERED'],
        ['approach_phases.py:237', 'FinalPhaseState', 'ILS takeover, weather, ownership, control, throttle, flaps/gear', 'test_control_ownership, test_loc_approach, test_safety_guard'],
        ['approach_phases.py:608', 'LandingPhaseState', 'Flare, throttle management, touchdown detection', 'NOT COVERED'],
    ]
    for row_data in phase_rows:
        add_table_row(table, row_data)

    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════
    # 9. TELEMETRY RECORDER
    # ══════════════════════════════════════════════════════════════
    doc.add_heading('9. Telemetry Recorder', level=1)

    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, text in enumerate(['File:Line', 'Class/Function', 'Behavior', 'Test Coverage']):
        hdr[i].text = ''
        p = hdr[i].paragraphs[0]
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(9)
        run.font.name = 'Consolas'
        set_cell_shading(hdr[i], '1F3A5F')
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    recorder_rows = [
        ['telemetry_recorder.py:105', 'TelemetryRecorder', 'CSV append-only, pre-defined schema, read-only', 'test_telemetry_recorder (22)'],
        ['telemetry_recorder.py:205', 'set_pending_frame()', 'Buffer terminal frame, flush AFTER actuators', 'VERIFIED'],
        ['telemetry_recorder.py:222', 'flush_pending_frame()', 'Write pending to disk, never raises', 'VERIFIED'],
    ]
    for row_data in recorder_rows:
        add_table_row(table, row_data)

    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════
    # 10. OTHER MODULES
    # ══════════════════════════════════════════════════════════════
    doc.add_heading('10. Other Modules', level=1)

    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, text in enumerate(['Module', 'Class/Function', 'Behavior', 'Test Coverage']):
        hdr[i].text = ''
        p = hdr[i].paragraphs[0]
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(9)
        run.font.name = 'Consolas'
        set_cell_shading(hdr[i], '1F3A5F')
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    other_rows = [
        ['synthetic_glidepath.py', 'SyntheticGlidepath', 'VOR/NDB/LOC: geometry-based VS, MDA floor', 'test_synthetic_glidepath (19)'],
        ['wind_correction.py', 'WindCorrection', 'Wind correction for heading and VS', 'test_loc_approach (pipeline)'],
        ['ils_navigation.py', 'ILSNavigation', 'ILS/LOC: localizer deviation, glideslope', 'test_loc_approach (25)'],
        ['engine_failure_detector.py', 'EngineFailureDetector', 'N1/EGT monitoring, asymmetric thrust', 'test_engine_failure (15)'],
        ['approach_speed_calculator.py', 'ApproachSpeedCalculator', 'VREF/VAPP calculation', 'NOT COVERED'],
        ['stabilized_approach.py', 'StabilizedApproachMonitor', 'Stabilized criteria check', 'NOT COVERED'],
        ['flare_controller.py', 'FlareController', 'Flare start, target pitch/VS', 'NOT COVERED'],
        ['dme_navigation.py', 'DMENavigation', 'DME fix tracking, altitude checks', 'NOT COVERED'],
        ['navigation.py', 'Navigation', 'VOR geometry, distance/bearing', 'NOT COVERED'],
        ['virtual_joystick.py', 'VirtualJoystick', 'vJoy device control', 'NOT COVERED'],
        ['connection_monitor.py', 'ConnectionMonitor', 'Method switching, metrics', 'NOT COVERED'],
        ['connection_optimizer.py', 'ConnectionOptimizer', 'L:Vars vs WASM vs SimConnect', 'NOT COVERED'],
        ['autothrottle.py', 'AutothrottleController', 'PID throttle, asymmetric mode', 'NOT COVERED'],
        ['aircraft_adapter.py', 'AircraftCommandAdapter', 'AP disengage via L:Vars/WASM', 'NOT COVERED (readback only)'],
        ['audio_alerts.py', 'AudioAlerts', 'gtts+pygame alerts', 'NOT COVERED'],
        ['wind_shear_detector.py', 'WindShearDetector', 'Wind shear detection', 'NOT COVERED'],
        ['turbulence_detector.py', 'TurbulenceDetector', 'Turbulence from G-force', 'NOT COVERED'],
        ['structured_logger.py', 'StructuredLogger', 'Category logging + SQLite', 'NOT COVERED'],
        ['settings.py', 'Settings', 'JSON config', 'NOT COVERED'],
        ['aircraft_config_reader.py', 'AircraftConfigReader', 'JSON aircraft profiles', 'NOT COVERED'],
        ['airports_database.py', 'AirportsDatabase', 'JSON airport data', 'NOT COVERED'],
        ['fms_reader.py', 'FMSReader', 'FMS waypoint data', 'NOT COVERED'],
        ['navigraph_parser.py', 'NavigraphParser', 'Chart parsing', 'NOT COVERED'],
        ['wasm_interface.py', 'WASMInterface', 'WASM L:Var read/write', 'NOT COVERED'],
        ['aileron_compensation.py', 'AileronCompensation', 'Engine failure aileron', 'NOT COVERED'],
        ['rudder_compensation.py', 'RudderCompensation', 'Engine failure rudder', 'NOT COVERED'],
        ['approach_dialog.py', 'ApproachDialog', 'GUI dialog (tkinter)', 'NOT COVERED'],
        ['settings_dialog.py', 'SettingsDialog', 'GUI settings', 'NOT COVERED'],
        ['gui.py', 'GUI application', 'tkinter main window', 'NOT COVERED'],
    ]
    for row_data in other_rows:
        add_table_row(table, row_data)

    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════
    # 11. VERIFIED FACTS
    # ══════════════════════════════════════════════════════════════
    doc.add_heading('11. Verified Facts', level=1)

    verified = [
        'HEAD == master == origin/master == 23599b6 (trusted baseline)',
        'Tree SHA == f2a19ec8 (expected)',
        '208 tests, 0 failed',
        'Python 3.14.5, pytest 9.0.3, SimConnect 0.4.26',
        'Safety guard G1-G5 rules verified by unit tests',
        'Readback-verified takeover verified by unit tests',
        'Ownership matrix verified by unit tests',
        'Telemetry recorder pending-frame-after-actuator verified',
        'LOC signal loss -> go-around verified',
        'State pattern architecture: CC reduced from 76 to 5',
        'Guard runs BEFORE phase_state.handle() in FINAL only',
        'Per-rule debounce N=2 with reset on clean frame',
        'Fail-closed: unknown mode -> no command, None telemetry -> guard fires',
    ]
    for item in verified:
        p = doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════
    # 12. UNVERIFIED CLAIMS
    # ══════════════════════════════════════════════════════════════
    doc.add_heading('12. Unverified Claims', level=1)

    unverified = [
        'ConnectionOptimizer.test_all_methods() performance recommendations',
        'ConnectionMonitor.should_switch_method() switching logic',
        'AutothrottleController.calculate_throttle() PID behavior',
        'FlareController.calculate_flare_parameters() flare math',
        'StabilizedApproachMonitor.should_go_around() stabilization criteria',
        'DMENavigation.check_altitude_at_fix() fix checking',
        'Navigation.calculate_vor_approach() VOR geometry',
        'SyntheticGlidepath.compute_target_vs() through main.py integration',
        'Audio alert pregeneration and playback',
        'GUI (tkinter) functionality',
        'Settings persistence (JSON)',
    ]
    for item in unverified:
        p = doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════
    # 13. ENVIRONMENT-DEPENDENT AREAS
    # ══════════════════════════════════════════════════════════════
    doc.add_heading('13. Environment-Dependent Areas', level=1)

    env_dep = [
        'MSFSTelemetry.connect() — requires running MSFS + SimConnect',
        'MSFSControl.* — requires SimConnect connection',
        'VirtualJoystick.* — requires vJoy driver',
        'AircraftCommandAdapter.detect_and_configure() — requires MSFS aircraft detection',
        'ConnectionOptimizer.test_all_methods() — requires live MSFS',
        'AudioAlerts.pregenerate_alerts() — requires gtts + pygame',
        'gui.py — requires tkinter display',
    ]
    for item in env_dep:
        p = doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════
    # 14. OPEN QUESTIONS
    # ══════════════════════════════════════════════════════════════
    doc.add_heading('14. Open Questions', level=1)

    questions = [
        '2 deleted task files (tasks/TASK-002, tasks/TASK-006) in working tree — intentional cleanup or accidental?',
        'tasks/reports/TASK-002-stage0-safety-core-report.md still tracked — should it be removed too?',
        'Root-level test files not in tests/ directory — intentional organization or inconsistency?',
        'test_mobiflight_wasm.py collected 1 test but root collection errored on pytest cleanup — functional?',
    ]
    for i, item in enumerate(questions, 1):
        p = doc.add_paragraph(f'{i}. {item}')

    # ══════════════════════════════════════════════════════════════
    # SAVE
    # ══════════════════════════════════════════════════════════════
    output_path = r'C:\BAT\msfs_autoland\TASKS\reports\reconstruction_report_23599b6.docx'
    doc.save(output_path)
    print(f'Report saved: {output_path}')


if __name__ == '__main__':
    create_report()
